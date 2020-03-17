from docx import Document
from docx.oxml.ns import qn
# 文档对象
document = Document()
document.styles['Normal'].font.name='宋体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
# 添加标题
document.add_heading('标题1',level=2)

# 段落
p = document.add_paragraph('段落1')

# 加粗:
p.add_run('加粗').bold=True
p.add_run('普通')
#斜体
p.add_run('斜体').italic=True

# 新起一段
p2 = document.add_paragraph('新段落')


# 列表
document.add_paragraph(
    '无序列表',style='List Bullet'
)
document.add_paragraph(
    '有序列表',style='List Number'
)
document.save('word.docx')