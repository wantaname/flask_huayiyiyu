from db import database_conn
from docx import Document
from docx.oxml.ns import qn
import pymysql

def word():
    document = Document()
    document.styles['Normal'].font.name = '宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    conn = database_conn()
    # 配置结果集为字典形式
    cursor = conn.cursor(cursor=pymysql.cursors.DictCursor)
    sql='select clothes from clothes'
    cursor.execute(sql)
    res=cursor.fetchall()

    clothes = []
    for item in res:
        clothes.append(item['clothes'])
    print(clothes)
    clothes=list(set(clothes))
    print(clothes)
    for item in clothes:

        sql = 'select category from clothes WHERE clothes="{}"'.format(item)
        cursor.execute(sql)
        res = cursor.fetchall()

        for cate in res:
            cate_item = cate['category']
            sql = 'select * from clothes WHERE clothes="{}" and category="{}"'.format(item,cate_item)
            cursor.execute(sql)
            res = cursor.fetchone()
            print(res)

def word2():
    document = Document()
    document.styles['Normal'].font.name = '宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    conn = database_conn()
    # 配置结果集为字典形式
    cursor = conn.cursor(cursor=pymysql.cursors.DictCursor)
    sql = 'select * from clothes'
    cursor.execute(sql)
    res = cursor.fetchall()
    # item为字典
    clothes=''
    category=''
    for item in res:
        if clothes != item['clothes']:
            clothes = item['clothes']
            #打印标题
            document.add_heading(clothes, level=1)

        category = item['category']
        #打印条目
        document.add_heading(category, level=4)
        #打印书: 有内容才打印
        item.pop('Id')
        item.pop('clothes')
        item.pop('category')
        for i in item.keys():
            #  如果有值
            if item[i]:
                sql = 'select name from type WHERE code="{}"'.format(i[3:])
                cursor.execute(sql)
                res = cursor.fetchone()
                key=res['name']
                value=item[i]

                document.add_paragraph(
                    '{}: {}'.format(key,value), style='List Bullet'
                )
    conn.close()
    document.save('word.docx')
